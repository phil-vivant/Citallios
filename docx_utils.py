from docx import Document
from docx.shared import Cm, Pt, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from PIL import Image  # pip install pillow



def set_margins(section, top=None, right=None, bottom=None, left=None, unit="cm"):
    conv = {"cm": Cm, "mm": Mm}[unit]
    if top    is not None: section.top_margin    = conv(top)
    if right  is not None: section.right_margin  = conv(right)
    if bottom is not None: section.bottom_margin = conv(bottom)
    if left   is not None: section.left_margin   = conv(left)


def _fit_width_for_box(image_path: str, max_width_cm: float, max_height_cm: float) -> float:
    """
    Calcule la largeur (en cm) à utiliser pour insérer l'image afin de respecter
    simultanément max_width_cm et max_height_cm, en conservant le ratio.
    On spécifie seulement 'width' à python-docx; la hauteur s'ajuste.
    """
    with Image.open(image_path) as im:
        w, h = im.size  # pixels
    # Si on fixe la largeur à W, la hauteur devient H = W * (h/w)
    # On veut H <= max_height_cm -> W <= max_height_cm * (w/h)
    w_limit_by_height = max_height_cm * (w / h)
    target_w_cm = min(max_width_cm, w_limit_by_height)
    return target_w_cm


def add_image_and_caption_below(
        table, row_img: int, col: int, image_path: str, caption: str,
        max_width_cm: float = 7.0, max_height_cm: float = 7.0,
        caption_size_pt: int = 9, italic: bool = True
):
    """
    Met l'image en (row_img, col), la légende en (row_img+1, col).
    Crée les paragraphes et centre image + légende.
    Redimensionne l'image pour respecter max_width ET max_height.
    """
    # Sécurité: s'assurer que la table a assez de lignes
    needed_rows = row_img + 2
    while len(table.rows) < needed_rows:
        table.add_row()

    # --- cellule image ---
    cell_img = table.cell(row_img, col)
    # nettoyer le contenu initial si nécessaire
    if not cell_img.paragraphs:
        p_img = cell_img.add_paragraph()
    else:
        p_img = cell_img.paragraphs[0]
        p_img.clear() if hasattr(p_img, "clear") else None
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Taille ajustée
    target_w_cm = _fit_width_for_box(image_path, max_width_cm, max_height_cm)

    run_img = p_img.add_run()
    run_img.add_picture(image_path, width=Cm(target_w_cm))
    cell_img.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # --- cellule légende (en dessous) ---
    cell_cap = table.cell(row_img + 1, col)
    p_cap = cell_cap.paragraphs[0] if cell_cap.paragraphs else cell_cap.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_cap.add_run(caption)
    r.font.size = Pt(caption_size_pt)
    r.font.italic = italic
    cell_cap.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


# --- helpers ---
def _add_symbol_with_sub(cell, base: str, sub: str):
    p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    # pas de p.clear() (non public) → on ajoute simplement
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.add_run(base)
    r = p.add_run(sub)
    r.font.subscript = True
    return p

def _add_text(cell, text, align=WD_ALIGN_PARAGRAPH.RIGHT, bold=False, size=11):
    p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    return p

def _fmt(x, nd=1):
    return f"{x:.{nd}f}"

def make_sls_table(container, sls_results: dict, *, style: str | None = None):
    """
    Ajoute un tableau SLS dans `container` (Document ou cellule de tableau) et le retourne.
    Lignes : Moment / σ_c / σ_s / σ_r / σ_f
    Colonnes : Phase 1 / Phase 2 / Bilan (avec une colonne de labels à gauche).

    sls_results clés attendues :
      sigma_c1, sigma_c2, sigma_s1, sigma_s2, sigma_sr2, sigma_f2, m1, m2
    """
    # valeurs
    c1 = sls_results["sigma_c1"]; c2 = sls_results["sigma_c2"]
    s1 = sls_results["sigma_s1"]; s2 = sls_results["sigma_s2"]
    r2 = sls_results["sigma_sr2"]; f2 = sls_results["sigma_f2"]
    m1 = sls_results["m1"];        m2 = sls_results["m2"]

    # bilans
    c = c1 + c2
    s = s1 + s2
    r = r2
    f = f2
    m12 = m1 + m2

    # 1 ligne d’en-tête + 5 lignes données ; 1 col labels + 3 colonnes phases
    table = container.add_table(rows=1 + 5, cols=1 + 3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    if style:
        table.style = style  # ex: "Light Shading Accent 1" (nom exact !)

    # En-têtes
    _add_text(table.cell(0,0), "", align=WD_ALIGN_PARAGRAPH.LEFT, bold=True)
    _add_text(table.cell(0,1), "Phase 1", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    _add_text(table.cell(0,2), "Phase 2", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    _add_text(table.cell(0,3), "Bilan",   align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)

    # Labels lignes
    labels = [
        "Moment sollicitant",
        "Contrainte béton",
        "Contrainte acier",
        "Contrainte acier renf",
        "Contrainte carbone",
    ]
    for i, lab in enumerate(labels, start=1):
        _add_text(table.cell(i,0), lab, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)

    # Ligne Moment (kN.m)
    _add_symbol_with_sub(table.cell(1,1), "M", "1");  _add_text(table.cell(1,1), f" = {_fmt(m1,1)} kN.m")
    _add_symbol_with_sub(table.cell(1,2), "M", "2");  _add_text(table.cell(1,2), f" = {_fmt(m2,1)} kN.m")
    _add_symbol_with_sub(table.cell(1,3), "M", "12"); _add_text(table.cell(1,3), f" = {_fmt(m12,1)} kN.m")

    # Béton (MPa)
    _add_symbol_with_sub(table.cell(2,1), "σ", "c1"); _add_text(table.cell(2,1), f" = {_fmt(c1,2)} MPa")
    _add_symbol_with_sub(table.cell(2,2), "σ", "c2"); _add_text(table.cell(2,2), f" = {_fmt(c2,2)} MPa")
    _add_symbol_with_sub(table.cell(2,3), "σ", "c");  _add_text(table.cell(2,3), f" = {_fmt(c,2)} MPa")

    # Acier (MPa)
    _add_symbol_with_sub(table.cell(3,1), "σ", "s1"); _add_text(table.cell(3,1), f" = {_fmt(s1,1)} MPa")
    _add_symbol_with_sub(table.cell(3,2), "σ", "s2"); _add_text(table.cell(3,2), f" = {_fmt(s2,1)} MPa")
    _add_symbol_with_sub(table.cell(3,3), "σ", "s");  _add_text(table.cell(3,3), f" = {_fmt(s,1)} MPa")

    # Acier renfort (MPa)
    _add_symbol_with_sub(table.cell(4,1), "σ", "r1"); _add_text(table.cell(4,1), f" = {_fmt(0.0,1)} MPa")
    _add_symbol_with_sub(table.cell(4,2), "σ", "r2"); _add_text(table.cell(4,2), f" = {_fmt(r2,1)} MPa")
    _add_symbol_with_sub(table.cell(4,3), "σ", "r");  _add_text(table.cell(4,3), f" = {_fmt(r,1)} MPa")

    # Carbone (MPa)
    _add_symbol_with_sub(table.cell(5,1), "σ", "f1"); _add_text(table.cell(5,1), f" = {_fmt(0.0,1)} MPa")
    _add_symbol_with_sub(table.cell(5,2), "σ", "f2"); _add_text(table.cell(5,2), f" = {_fmt(f2,1)} MPa")
    _add_symbol_with_sub(table.cell(5,3), "σ", "f");  _add_text(table.cell(5,3), f" = {_fmt(f,1)} MPa")

    # (Optionnel) largeur colonne labels
    for i, cell in enumerate(table.column_cells(0)):
        if i == 0:
            cell.width = Cm(5)
        else:
            cell.width = Cm(4)

    return table
