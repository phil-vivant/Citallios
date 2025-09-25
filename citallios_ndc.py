	
from docx import Document
from docx.shared import Inches, Cm
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT

from docx_utils import add_image_and_caption_below, make_sls_table

from section_flex.section.plane_of_deformation import PlaneOfDeformation


sls_results = {
    'm1': 30,
    'm2': 30,
    'sigma_c1': 9.436097151192582,
    'sigma_c2': 6.027386524556387,
    'sigma_s1': -333.0576837159865,
    'sigma_s2': -132.3331068970313,
    'sigma_sr2': -144.52108140703177,
    'sigma_f2': -181.31780948273575,
    'pod_1': PlaneOfDeformation(
        epsilon_0=-0.00203958885919076,
        omega_y=0.009357511015270686,
        omega_z=4.309306828840466e-20
    ),
    'pod_2': PlaneOfDeformation(
        epsilon_0=-0.0008241718612851625,
        omega_y=0.004062658170000146,
        omega_z=-1.2042673654497328e-20
    )
}


document = Document()

# Section courante (la plupart des docs n'en ont qu'une)
section = document.sections[0]

# Exemple : marges en centimètres
section.top_margin = Cm(1.5)
section.bottom_margin = Cm(1.5)
section.left_margin = Cm(2.0)
section.right_margin = Cm(2.0)

document.add_heading("Renforcement d'une section en béton armé", 0)

document.add_paragraph("Étude d'une section en béton armé renforcée par armatures complémentaires.")
p = document.add_paragraph('Les armatures de renfort peuvent être des ')
p.add_run('barres HA').bold = True
p.add_run(' ou des ')
p.add_run('lamelles Carbones.').bold = True

document.add_heading('Hypothèses', level=1)

document.add_paragraph('Géométrie et Matériaux', style='Intense Quote')
p = document.add_paragraph(f"\t\tSection :\t1,00 m x 0,25 m ht")
p.runs[-1].add_break()
p.add_run(f"\t\tBéton :\t\tfck = 25 MPa")
p.runs[-1].add_break()
p.add_run(f"\t\tAcier HA :\tfyk = 500 MPa\tClasse 'B'")
p.runs[-1].add_break()
p.add_run(f"\t\tCarbone :\tEf = 220 GPa\tSigma_sls = 1200 MPa\tSigma_uls = 1500 MPa")

document.add_paragraph('Armatures', style='Intense Quote')
p = document.add_paragraph(f"\t\tBarres HA :\t4 x As = 1.13 cm²\td's = 0.040 m")
p.runs[-1].add_break()
p.add_run(f"\t\tRenforts HA :\t2 x Ar = 1.13 cm²\td'r = 0.025 m")
p.runs[-1].add_break()
p.add_run(f"\t\tRenforts FRP :\t3 x Af = 0.91 cm²\td'f = 0.000 m")

document.add_paragraph('Sollicitations', style='Intense Quote')
p = document.add_paragraph(f"\t\tMoment ELS :\t\tM_1 = 30.0 kN.m\tM_2 = 30.0 kN.m")
p.runs[-1].add_break()
p.add_run(f"\t\tMoment ELU :\t\tM_ELU = 80.0 kN.m")
p.runs[-1].add_break()
p.add_run(f"\t\tMoment FEU :\t\tM_FEU = 60.0 kN.m")


document.add_heading('Géométrie', level=1)

table = document.add_table(rows=2, cols=2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.autofit = False
for col in table.columns:
    for cell in col.cells:
        cell.width = Cm(8)

# Remplir colonne 1
add_image_and_caption_below(
    table, row_img=0, col=0,
    image_path="img/monty_truth.png",
    caption="Section Avant Renforcement",
    max_width_cm=7, max_height_cm=7,
)
# Remplir colonne 2
add_image_and_caption_below(
    table, row_img=0, col=1,
    image_path="img/monty_truth.png",
    caption="Section Après Renforcement",
    max_width_cm=7, max_height_cm=7,
)

document.add_page_break()

document.add_heading("Vérification de la section à l'ELS", level=1)

document.add_paragraph('Tableau des contraintes élastiques - Calcul Phasé', style='Intense Quote')
tbl = make_sls_table(document, sls_results, style="Light Shading Accent 1"  # adapte au nom exact dispo chez toi
)


# document.add_paragraph("Phase 1 - État de contraintes ELS dans la section", style='Intense Quote')
# p = document.add_paragraph(f"\t\tMoment sollicitant :\tM_1  = 30.0 kN.m")
# p.runs[-1].add_break()
# p.add_run(f"\t\tContrainte béton :\tσ_c1 = 9.44 MPa")
# p.runs[-1].add_break()
# p.add_run(f"\t\tContrainte acier :\tσ_s1 = -333.1 MPa")
# p.runs[-1].add_break()
# p.add_run(f"\t\tContrainte acier renf :\tσ_r1 = 0.0 MPa")
# p.runs[-1].add_break()
# p.add_run(f"\t\tContrainte carbone :\tσ_f1 = 0.0 MPa")


# document.add_paragraph("Phase 2 - État de contraintes ELS dans la section", style='Intense Quote')
# p = document.add_paragraph(f"\t\tMoment sollicitant :\tM_2  = 30.0 kN.m")
# p.runs[-1].add_break()
# p.add_run(f"\t\tContrainte béton :\tσ_c2 = 6.03 MPa")
# p.runs[-1].add_break()
# p.add_run(f"\t\tContrainte acier :\tσ_s2 = -132.3 MPa")
# p.runs[-1].add_break()
# p.add_run(f"\t\tContrainte acier renf :\tσ_r2 = -144.5 MPa")
# p.runs[-1].add_break()
# p.add_run(f"\t\tContrainte carbone :\tσ_f2 = -181.3 MPa")


# document.add_paragraph("Bilan - État de contraintes ELS dans la section", style='Intense Quote')
# p = document.add_paragraph(f"\t\tMoment sollicitant :\tM_12  = 60.0 kN.m")
# p.runs[-1].add_break()
# p.add_run(f"\t\tContrainte béton :\tσ_c = 15.46 MPa")
# p.runs[-1].add_break()
# p.add_run(f"\t\tContrainte acier :\tσ_s = -465.4 MPa")
# p.runs[-1].add_break()
# p.add_run(f"\t\tContrainte acier renf :\tσ_r = -144.5 MPa")
# p.runs[-1].add_break()
# p.add_run(f"\t\tContrainte carbone :\tσ_f = -181.3 MPa")

document.add_page_break()


document.add_heading("Vérification de la section à l'ELU", level=1)

document.add_paragraph("Capacité résistante ELU", style='Intense Quote')
p = document.add_paragraph(f"\t\tMoment sollicitant :\tM_Ed  = 80.0 kN.m")
p.runs[-1].add_break()
p.add_run(f"\t\tAvant renforcement :\tM_Rd1 = 42.8 kN.m")
p.runs[-1].add_break()
p.add_run(f"\t\tAprès renforcement :\tM_Rd2 = 143.4 kN.m")

document.add_paragraph("Équilibre de la section renforcée à l'ELU", style='Intense Quote')
p = document.add_paragraph(f"\t\tContrainte béton:\tσ_c = 10.92 MPa")
p.runs[-1].add_break()
p.add_run(f"\t\tContrainte acier:\tσ_s = -360.8 MPa")
p.runs[-1].add_break()
p.add_run(f"\t\tContrainte acier renf:\tσ_r = -398.3 MPa")
p.runs[-1].add_break()
p.add_run(f"\t\tContrainte carbone:\tσ_f = -507.0 MPa")


table = document.add_table(rows=2, cols=1)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
# table.autofit = False
# for col in table.columns:
#     for cell in col.cells:
#         cell.width = Cm(8)

# Remplir colonne 1
add_image_and_caption_below(
    table, row_img=0, col=0,
    image_path="img/monty_truth.png",
    caption="Diagramme d'interaction ELU de la section (renforcée et non renforcée)",
    max_width_cm=12, max_height_cm=10,
)

document.add_page_break()


document.add_heading("Vérification de la section en situation d'incendie", level=1)

document.add_paragraph("Capacité résistante au fEU", style='Intense Quote')
p = document.add_paragraph(f"\t\tMoment sollicitant :\tM_Ed  = 60.0 kN.m")
p.runs[-1].add_break()
p.add_run(f"\t\tAvant renforcement :\tM_Rd1 = 49.6 kN.m")
p.runs[-1].add_break()
p.add_run(f"\t\tAprès renforcement :\tM_Rd2 = 74.6 kN.m")

document.add_paragraph("Équilibre de la section renforcée au feu", style='Intense Quote')
p = document.add_paragraph(f"\t\tContrainte béton:\tσ_c = 12.57 MPa")
p.runs[-1].add_break()
p.add_run(f"\t\tContrainte acier:\tσ_s = -430.0 MPa")
p.runs[-1].add_break()
p.add_run(f"\t\tContrainte acier renf:\tσ_r = -469.2 MPa")
p.runs[-1].add_break()
p.add_run(f"\t\tContrainte carbone:\tσ_f = -0.0 MPa")


table = document.add_table(rows=2, cols=1)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
# table.autofit = False
# for col in table.columns:
#     for cell in col.cells:
#         cell.width = Cm(8)

# Remplir colonne 1
add_image_and_caption_below(
    table, row_img=0, col=0,
    image_path="img/monty_truth.png",
    caption="Diagramme d'interaction au feu de la section (renforcée et non renforcée)",
    max_width_cm=12, max_height_cm=10,
)

document.add_page_break()


document.save('ndc.docx')

# for s in document.styles:
#     if s.type == WD_STYLE_TYPE.TABLE:
#         print(s.name)
